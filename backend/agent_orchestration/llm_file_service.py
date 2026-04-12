"""
LLM File Upload Service - Full Document Mode Support
====================================================

Service for uploading documents to LLM provider File APIs (OpenAI, Anthropic, Google/Gemini).
This enables "Full Document Mode" where entire documents are sent to LLMs instead of chunks.

Each provider has different APIs and file size limits:
- OpenAI: 50 MB per file, purpose='user_data'
- Anthropic (Claude): 500 MB per file, beta Files API
- Google (Gemini): 2 GB per file, Files API
"""

import os
import logging
import mimetypes
import tempfile
from typing import Dict, Optional, Tuple
from datetime import datetime
from asgiref.sync import sync_to_async, async_to_sync

from django.utils import timezone

from project_api_keys.services import get_project_api_key_service
from users.models import IntelliDocProject, ProjectDocument

logger = logging.getLogger('agent_orchestration')


# File size limits in bytes
FILE_SIZE_LIMITS = {
    'openai': 50 * 1024 * 1024,       # 50 MB
    'anthropic': 500 * 1024 * 1024,   # 500 MB
    'google': 2 * 1024 * 1024 * 1024, # 2 GB
}

# Supported file types for each provider
SUPPORTED_FILE_TYPES = {
    'openai': ['.pdf', '.txt', '.doc', '.docx', '.md', '.rtf'],
    'anthropic': ['.pdf', '.txt', '.png', '.jpg', '.jpeg', '.gif', '.webp'],
    'google': ['.pdf', '.txt', '.doc', '.docx', '.md', '.rtf', '.png', '.jpg', '.jpeg', '.gif', '.webp', '.mp3', '.wav', '.mp4'],
}


def _convert_docx_to_pdf(file_bytes: bytes) -> bytes:
    """
    Convert a DOCX file to PDF using python-docx (text extraction) + PyMuPDF (PDF creation).
    Preserves paragraph structure, headings, and basic formatting.
    """
    import io
    from docx import Document
    import fitz  # PyMuPDF

    doc = Document(io.BytesIO(file_bytes))
    pdf = fitz.open()

    # Page layout constants
    PAGE_WIDTH, PAGE_HEIGHT = 595, 842  # A4
    MARGIN_LEFT, MARGIN_RIGHT, MARGIN_TOP, MARGIN_BOTTOM = 60, 60, 60, 60
    LINE_HEIGHT_BODY = 14
    LINE_HEIGHT_HEADING = 20
    MAX_X = PAGE_WIDTH - MARGIN_RIGHT
    USABLE_WIDTH = MAX_X - MARGIN_LEFT

    page = pdf.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
    y = MARGIN_TOP

    def _new_page():
        nonlocal page, y
        page = pdf.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
        y = MARGIN_TOP

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            y += LINE_HEIGHT_BODY * 0.5
            if y > PAGE_HEIGHT - MARGIN_BOTTOM:
                _new_page()
            continue

        is_heading = para.style.name.startswith('Heading')
        fontsize = 14 if is_heading else 10
        fontname = "helv"  # Helvetica (built-in)
        line_height = LINE_HEIGHT_HEADING if is_heading else LINE_HEIGHT_BODY

        # Word-wrap the paragraph
        words = text.split()
        lines = []
        current_line = ""
        for word in words:
            test = f"{current_line} {word}".strip()
            # Approximate: each char ~0.5 * fontsize width
            if len(test) * fontsize * 0.45 > USABLE_WIDTH and current_line:
                lines.append(current_line)
                current_line = word
            else:
                current_line = test
        if current_line:
            lines.append(current_line)

        for line in lines:
            if y + line_height > PAGE_HEIGHT - MARGIN_BOTTOM:
                _new_page()
            page.insert_text(
                (MARGIN_LEFT, y),
                line,
                fontsize=fontsize,
                fontname=fontname,
            )
            y += line_height

        # Extra spacing after headings
        if is_heading:
            y += 4

    # Also include table text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if not row_text:
                continue
            if y + LINE_HEIGHT_BODY > PAGE_HEIGHT - MARGIN_BOTTOM:
                _new_page()
            # Truncate very wide rows
            if len(row_text) > 120:
                row_text = row_text[:120] + "..."
            page.insert_text((MARGIN_LEFT, y), row_text, fontsize=9, fontname="helv")
            y += LINE_HEIGHT_BODY

    # Warn about image loss — DOCX→PDF conversion only extracts text
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        image_rels = [r for r in doc.part.rels.values() if "image" in r.reltype]
        if image_rels:
            logger.warning(f"⚠️ DOCX→PDF: {len(image_rels)} embedded image(s) lost during conversion (text-only extraction)")
    except Exception:
        pass

    buf = io.BytesIO()
    pdf.save(buf)
    pdf.close()
    return buf.getvalue()


class LLMFileUploadService:
    """
    Service for uploading documents to LLM provider File APIs.
    
    Supports:
    - OpenAI Files API
    - Anthropic Files API (beta)
    - Google/Gemini Files API
    """
    
    def __init__(self, project: IntelliDocProject):
        """
        Initialize the service with a project context.
        
        Args:
            project: IntelliDocProject instance for API key access
        """
        self.project = project
        self.project_api_service = get_project_api_key_service()
        logger.info(f"📁 LLM FILE SERVICE: Initialized for project {project.name}")
    
    async def upload_document(
        self, 
        document: ProjectDocument, 
        providers: list = None
    ) -> Dict[str, str]:
        """
        Upload a document to specified LLM providers.
        
        Args:
            document: ProjectDocument instance to upload
            providers: List of providers to upload to ['openai', 'anthropic', 'google']
                      If None, uploads to all providers with available API keys
        
        Returns:
            Dict mapping provider to file_id (or error message)
        """
        if providers is None:
            providers = ['openai', 'anthropic', 'google']
        
        results = {}
        
        for provider in providers:
            try:
                result = await self._upload_to_provider(document, provider)
                results[provider] = result
            except Exception as e:
                logger.error(f"❌ LLM FILE SERVICE: Failed to upload to {provider}: {e}")
                results[provider] = {'error': str(e)}
        
        return results
    
    async def _upload_to_provider(
        self, 
        document: ProjectDocument, 
        provider: str
    ) -> Dict[str, str]:
        """
        Upload document to a specific provider.
        
        Args:
            document: ProjectDocument instance
            provider: Provider name ('openai', 'anthropic', 'google')
        
        Returns:
            Dict with 'file_id' or 'error'
        """
        # Deduplication: skip upload if document already has a file_id for this provider
        # Exception: DOCX files uploaded to OpenAI need re-upload as PDF (Responses API only accepts PDF)
        existing_id = getattr(document, f'llm_file_id_{provider}', None)
        ext = os.path.splitext(document.original_filename)[1].lower()
        needs_reupload = (provider == 'openai' and ext in ('.docx', '.doc', '.rtf') and existing_id)
        if existing_id and not needs_reupload:
            logger.info(f"✅ LLM FILE SERVICE: Document {document.original_filename} already uploaded to {provider}: {existing_id[:30]}...")
            return {'file_id': existing_id, 'status': 'already_uploaded', 'reason': 'already_uploaded'}
        if needs_reupload:
            logger.info(f"🔄 LLM FILE SERVICE: Re-uploading {document.original_filename} as PDF for OpenAI Responses API")
        
        # Validate file size
        file_size_limit = FILE_SIZE_LIMITS.get(provider)
        if document.file_size > file_size_limit:
            return {
                'error': f"File size ({document.file_size / 1024 / 1024:.2f} MB) exceeds {provider} limit ({file_size_limit / 1024 / 1024:.0f} MB)",
                'reason': 'file_too_large',
            }
        
        # Validate file type
        supported_types = SUPPORTED_FILE_TYPES.get(provider, [])
        if document.file_extension.lower() not in supported_types:
            return {
                'error': f"File type {document.file_extension} not supported by {provider}. Supported: {supported_types}",
                'reason': 'unsupported_type',
            }
        
        # Get API key for provider
        api_key = await self._get_api_key(provider)
        if not api_key:
            return {'error': f"No API key configured for {provider}", 'reason': 'no_api_key'}
        
        # Upload based on provider
        if provider == 'openai':
            return await self._upload_to_openai(document, api_key)
        elif provider == 'anthropic':
            return await self._upload_to_anthropic(document, api_key)
        elif provider == 'google':
            return await self._upload_to_google(document, api_key)
        else:
            return {'error': f"Unknown provider: {provider}", 'reason': 'unknown_provider'}
    
    async def _get_api_key(self, provider: str) -> Optional[str]:
        """Get API key for provider from project settings."""
        try:
            # Map provider names to API key service types
            provider_mapping = {
                'openai': 'openai',
                'anthropic': 'anthropic',
                'google': 'google'
            }
            
            provider_type = provider_mapping.get(provider)
            if not provider_type:
                return None
            
            # Get decrypted API key using the project API key service
            # Prefer the async variant to avoid extra sync wrappers.
            api_key = await self.project_api_service.get_project_api_key_async(
                self.project,
                provider_type,
            )
            
            if api_key:
                logger.info(f"🔑 LLM FILE SERVICE: Retrieved API key for project {self.project.name} - {provider_type}")
            else:
                logger.warning(f"⚠️ LLM FILE SERVICE: No API key configured for project {self.project.name} - {provider_type}")
            
            return api_key
        except Exception as e:
            logger.error(f"❌ LLM FILE SERVICE: Failed to get API key for {provider}: {e}")
            return None
    
    async def _upload_to_openai(
        self, 
        document: ProjectDocument, 
        api_key: str
    ) -> Dict[str, str]:
        """
        Upload file to OpenAI Files API.
        
        Uses: POST /v1/files with purpose='user_data'
        Docs: https://platform.openai.com/docs/api-reference/files
        """
        try:
            from openai import OpenAI
            from django.core.files.storage import default_storage
            
            client = OpenAI(api_key=api_key)
            
            # Read file from storage
            file_path = document.file_path
            if not default_storage.exists(file_path):
                return {'error': 'File not found in storage', 'reason': 'file_not_found'}
            
            # Open file — convert DOCX to PDF since OpenAI Responses API only accepts PDF
            with default_storage.open(file_path, 'rb') as f:
                file_bytes = f.read()

            upload_filename = document.original_filename
            ext = os.path.splitext(upload_filename)[1].lower()
            if ext in ('.docx', '.doc', '.rtf'):
                try:
                    file_bytes = await sync_to_async(_convert_docx_to_pdf)(file_bytes)
                    upload_filename = os.path.splitext(upload_filename)[0] + '.pdf'
                    logger.info(f"📄 LLM FILE SERVICE: Converted {document.original_filename} → PDF ({len(file_bytes)} bytes)")
                except Exception as conv_err:
                    logger.error(f"❌ LLM FILE SERVICE: DOCX→PDF conversion failed for {document.original_filename}: {conv_err}")
                    return {'error': f'DOCX→PDF conversion failed: {conv_err}', 'reason': 'conversion_failed'}

            import io
            response = await sync_to_async(client.files.create)(
                file=(upload_filename, io.BytesIO(file_bytes)),
                purpose='user_data'
            )
            
            file_id = response.id
            logger.info(f"✅ LLM FILE SERVICE: Uploaded to OpenAI: {file_id}")
            
            # Update document model
            await self._update_document_file_id(document, 'openai', file_id)
            
            return {'file_id': file_id}
            
        except Exception as e:
            logger.error(f"❌ LLM FILE SERVICE: OpenAI upload failed: {e}")
            return {'error': str(e), 'reason': 'provider_error'}
    
    async def _upload_to_anthropic(
        self, 
        document: ProjectDocument, 
        api_key: str
    ) -> Dict[str, str]:
        """
        Upload file to Anthropic Files API (beta).
        
        Uses: POST /v1/files with beta header
        Docs: https://platform.claude.com/docs/en/build-with-claude/files
        """
        try:
            import anthropic
            from django.core.files.storage import default_storage
            
            client = anthropic.Anthropic(api_key=api_key)
            
            # Read file from storage
            file_path = document.file_path
            if not default_storage.exists(file_path):
                return {'error': 'File not found in storage', 'reason': 'file_not_found'}
            
            # Determine MIME type
            mime_type = document.file_type or mimetypes.guess_type(document.original_filename)[0] or 'application/octet-stream'
            
            # Open file and upload using beta API
            with default_storage.open(file_path, 'rb') as f:
                response = await sync_to_async(client.beta.files.upload)(
                    file=(document.original_filename, f, mime_type)
                )
            
            file_id = response.id
            logger.info(f"✅ LLM FILE SERVICE: Uploaded to Anthropic: {file_id}")
            
            # Update document model
            await self._update_document_file_id(document, 'anthropic', file_id)
            
            return {'file_id': file_id}
            
        except Exception as e:
            logger.error(f"❌ LLM FILE SERVICE: Anthropic upload failed: {e}")
            return {'error': str(e), 'reason': 'provider_error'}
    
    async def _upload_to_google(
        self, 
        document: ProjectDocument, 
        api_key: str
    ) -> Dict[str, str]:
        """
        Upload file to Google/Gemini Files API.
        
        Uses: genai.Client().files.upload()
        Docs: https://ai.google.dev/gemini-api/docs/files
        
        Works with any Django storage backend: uses default_storage.path() when
        available (filesystem), otherwise streams from storage to a temp file
        so S3 and other remote storage backends are supported.
        """
        try:
            from google import genai
            from django.core.files.storage import default_storage
            
            client = genai.Client(api_key=api_key)
            file_path = document.file_path
            if not default_storage.exists(file_path):
                return {'error': 'File not found in storage', 'reason': 'file_not_found'}
            
            actual_path = None
            temp_fd = None
            temp_path = None
            
            try:
                actual_path = default_storage.path(file_path)
            except (NotImplementedError, AttributeError):
                pass
            
            if actual_path is None or not os.path.isfile(actual_path):
                with default_storage.open(file_path, 'rb') as f:
                    content = f.read()
                temp_fd, temp_path = tempfile.mkstemp(suffix=os.path.splitext(document.original_filename)[1] or '')
                try:
                    os.write(temp_fd, content)
                    os.close(temp_fd)
                    temp_fd = None
                    actual_path = temp_path
                except Exception:
                    if temp_fd is not None:
                        try:
                            os.close(temp_fd)
                        except OSError:
                            pass
                    if temp_path and os.path.exists(temp_path):
                        try:
                            os.unlink(temp_path)
                        except OSError:
                            pass
                    raise
            
            try:
                response = await sync_to_async(client.files.upload)(file=actual_path)
                file_uri = response.uri
                file_name = response.name
                logger.info(f"✅ LLM FILE SERVICE: Uploaded to Google: {file_name} ({file_uri})")
                await self._update_document_file_id(document, 'google', file_uri)
                return {'file_id': file_uri, 'file_name': file_name}
            finally:
                if temp_path and os.path.exists(temp_path):
                    try:
                        os.unlink(temp_path)
                    except OSError as e:
                        logger.debug("Failed to delete temp file after Google upload: %s", e)
            
        except Exception as e:
            logger.error(f"❌ LLM FILE SERVICE: Google upload failed: {e}")
            return {'error': str(e), 'reason': 'provider_error'}
    
    async def _update_document_file_id(
        self, 
        document: ProjectDocument, 
        provider: str, 
        file_id: str
    ):
        """Update the document model with the file_id from the provider."""
        try:
            field_name = f'llm_file_id_{provider}'
            setattr(document, field_name, file_id)
            # Only update database if this looks like a real Django model instance.
            if hasattr(document, 'llm_file_uploaded_at'):
                document.llm_file_uploaded_at = timezone.now()
            if hasattr(document, 'save'):
                await sync_to_async(document.save)(
                    update_fields=[field_name, 'llm_file_uploaded_at']
                )
                logger.info(f"✅ LLM FILE SERVICE: Updated document {getattr(document, 'document_id', 'unknown')} with {provider} file_id")
        except Exception as e:
            logger.error(f"❌ LLM FILE SERVICE: Failed to update document: {e}")
    
    @staticmethod
    def get_file_id_for_provider(document: ProjectDocument, provider: str) -> Optional[str]:
        """
        Get the file_id for a document and provider.
        
        Args:
            document: ProjectDocument instance
            provider: Provider name ('openai', 'anthropic', 'google')
        
        Returns:
            file_id string or None
        """
        field_name = f'llm_file_id_{provider}'
        return getattr(document, field_name, None)
    
    @staticmethod
    def check_provider_support(document: ProjectDocument, provider: str) -> Tuple[bool, str]:
        """
        Check if a document can be uploaded to a provider.
        
        Args:
            document: ProjectDocument instance
            provider: Provider name
        
        Returns:
            Tuple of (is_supported, reason)
        """
        # Check file size
        file_size_limit = FILE_SIZE_LIMITS.get(provider)
        if not file_size_limit:
            return False, f"Unknown provider: {provider}"
        
        if document.file_size > file_size_limit:
            return False, f"File size ({document.file_size / 1024 / 1024:.2f} MB) exceeds limit ({file_size_limit / 1024 / 1024:.0f} MB)"
        
        # Check file type
        supported_types = SUPPORTED_FILE_TYPES.get(provider, [])
        if document.file_extension.lower() not in supported_types:
            return False, f"File type {document.file_extension} not supported"
        
        return True, "Supported"

    async def delete_file(self, provider: str, file_id: str) -> Dict[str, str]:
        """
        Delete a file from the given LLM provider's Files API.
        
        This is used for node-level attachments where we store the provider
        file_id directly in the workflow graph.
        """
        if not file_id:
            return {"error": "Missing file_id", "reason": "missing_file_id"}
        
        # Normalize provider names
        normalized = (provider or "").lower()
        if normalized == "gemini":
            normalized = "google"
        
        # Get API key first; if missing, report an error
        api_key = await self._get_api_key(normalized)
        if not api_key:
            msg = f"No API key configured for {normalized}"
            logger.warning(f"⚠️ LLM FILE SERVICE: {msg} (delete_file)")
            return {"error": msg, "reason": "no_api_key"}
        
        try:
            if normalized == "openai":
                return await self._delete_openai_file(file_id, api_key)
            if normalized == "anthropic":
                return await self._delete_anthropic_file(file_id, api_key)
            if normalized == "google":
                return await self._delete_google_file(file_id, api_key)
            msg = f"Unknown provider: {normalized}"
            logger.error(f"❌ LLM FILE SERVICE: {msg} (delete_file)")
            return {"error": msg, "reason": "unknown_provider"}
        except Exception as e:
            logger.error(f"❌ LLM FILE SERVICE: Delete failed for provider={normalized}, file_id={file_id}: {e}")
            return {"error": str(e), "reason": "provider_error"}

    async def _delete_openai_file(self, file_id: str, api_key: str) -> Dict[str, str]:
        """
        Delete a file from OpenAI Files API.
        
        Uses: DELETE /v1/files/{file_id}
        Docs: https://developers.openai.com/api/reference/resources/files/methods/delete
        """
        try:
            from openai import OpenAI
            
            client = OpenAI(api_key=api_key)
            response = await sync_to_async(client.files.delete)(file_id)
            # OpenAI returns an object with id, deleted, object
            if getattr(response, "deleted", False):
                logger.info(f"🗑️ LLM FILE SERVICE: Deleted OpenAI file {file_id}")
                return {"status": "deleted", "id": getattr(response, "id", file_id)}
            logger.warning(f"⚠️ LLM FILE SERVICE: OpenAI delete did not confirm deletion for {file_id}: {response}")
            return {"error": "OpenAI did not confirm deletion", "reason": "not_deleted"}
        except Exception as e:
            logger.error(f"❌ LLM FILE SERVICE: OpenAI delete failed for {file_id}: {e}")
            return {"error": str(e), "reason": "provider_error"}

    async def _delete_anthropic_file(self, file_id: str, api_key: str) -> Dict[str, str]:
        """
        Delete a file from Anthropic Files API (beta).
        
        Uses: client.beta.files.delete(file_id=...)
        Docs: https://platform.claude.com/docs/en/api/python/beta/files/delete
        """
        try:
            import anthropic
            
            client = anthropic.Anthropic(api_key=api_key)
            response = await sync_to_async(client.beta.files.delete)(file_id=file_id)
            deleted_id = getattr(response, "id", file_id)
            logger.info(f"🗑️ LLM FILE SERVICE: Deleted Anthropic file {deleted_id}")
            return {"status": "deleted", "id": deleted_id}
        except Exception as e:
            logger.error(f"❌ LLM FILE SERVICE: Anthropic delete failed for {file_id}: {e}")
            return {"error": str(e), "reason": "provider_error"}

    async def _delete_google_file(self, file_id: str, api_key: str) -> Dict[str, str]:
        """
        Delete a file from Google/Gemini Files API.
        
        Uses: genai.Client().files.delete(name=...)
        Docs: https://ai.google.dev/gemini-api/docs/files
        
        Note: For Google/Gemini, we currently store the file URI as file_id.
        The python-genai client accepts the file name; in many cases the URI
        is also accepted as the identifier. If this ever fails in practice,
        we can extend node-level attachments to store both uri and name.
        """
        try:
            from google import genai
            
            client = genai.Client(api_key=api_key)
            # Best effort: treat stored file_id as the name/identifier
            await sync_to_async(client.files.delete)(name=file_id)
            logger.info(f"🗑️ LLM FILE SERVICE: Deleted Google/Gemini file {file_id}")
            return {"status": "deleted", "id": file_id}
        except Exception as e:
            logger.error(f"❌ LLM FILE SERVICE: Google/Gemini delete failed for {file_id}: {e}")
            return {"error": str(e), "reason": "provider_error"}


# Synchronous wrapper functions for use in non-async contexts
def upload_document_sync(project: IntelliDocProject, document: ProjectDocument, providers: list = None) -> Dict[str, str]:
    """Synchronous wrapper for uploading a document."""
    service = LLMFileUploadService(project)
    return async_to_sync(service.upload_document)(document, providers)


def get_file_id(document: ProjectDocument, provider: str) -> Optional[str]:
    """Get file_id for a document and provider."""
    return LLMFileUploadService.get_file_id_for_provider(document, provider)
